import openai
from docx import Document
import pymongo
import os
import fitz  # PyMuPDF for PDF handling
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

class MongoDBHandler:
    def __init__(self):
        """
        Initialize MongoDBHandler with MongoDB credentials and collection name.
        """
        db_name = "file_embeddings"
        collection_name = "embeddings"
        mongo_uri = "mongodb://localhost:27017/"
        self.client = pymongo.MongoClient(mongo_uri)
        self.db = self.client[db_name]
        self.collection_name = collection_name

    def upsert_document(self, document, query):
        """
        Inserts or updates a document in the MongoDB collection based on the query.
        """
        collection = self.db[self.collection_name]
        collection.update_one(query, {'$set': document}, upsert=True)

    def retrieve_embeddings(self):
        """
        Retrieves all embeddings from the specified MongoDB collection.
        """
        collection = self.db[self.collection_name]
        return list(collection.find())


class DocumentEmbedder:
    def __init__(self):
        """
        Initialize with OpenAI API key and a MongoDB handler instance.
        """
        openai.api_key = openai_api_key = os.getenv('AS_OPENAI_API_KEY')
        self.mongo_handler = MongoDBHandler()

    def embed_text(self, text):
        """
        Generates an embedding for a given text using OpenAI API.
        """
        response = openai.Embedding.create(
            input=text,
            model="text-embedding-ada-002"
        )
        return response['data'][0]['embedding']

    def process_word_document(self, docx_path):
        """
        Processes a Word document, generates embeddings for the title and table rows,
        and stores the embeddings in MongoDB.
        """
        doc = Document(docx_path)
        title = doc.paragraphs[0].text  # Assuming the title is the first paragraph
        title_embedding = self.embed_text(title)

        # Process the table
        rows = []
        for table in doc.tables:
            for row in table.rows:
                designation = row.cells[0].text.strip()
                explanation = row.cells[1].text.strip()
                combined_text = f"{designation}: {explanation}"
                rows.append(combined_text)

        # Create embeddings for each row in the table
        row_embeddings = [self.embed_text(row) for row in rows]

        # Store the embeddings in MongoDB
        document_data = {
            "file_name": os.path.basename(docx_path),
            "title": title,
            "title_embedding": title_embedding,
            "row_embeddings": row_embeddings
        }

        self.mongo_handler.upsert_document(
            document=document_data,
            query={"file_name": os.path.basename(docx_path)}
        )

    def process_pdf_document(self, pdf_path):
        """
        Processes a PDF document, generates embeddings for the text,
        and stores the embeddings in MongoDB.
        """
        doc = fitz.open(pdf_path)
        text = ""
        for page in doc:
            text += page.get_text()

        # Embed the entire PDF content
        text_embedding = self.embed_text(text)

        document_data = {
            "file_name": os.path.basename(pdf_path),
            "text": text,
            "text_embedding": text_embedding
        }

        self.mongo_handler.upsert_document(
            document=document_data,
            query={"file_name": os.path.basename(pdf_path)}
        )

    def process_directory(self, directory):
        """
        Process all PDF and Word files in a directory.
        """
        for file_name in os.listdir(directory):
            file_path = os.path.join(directory, file_name)
            if file_name.endswith(".docx") and not file_name.startswith('~$'):
                try:
                    self.process_word_document(file_path)
                except Exception as e:
                    print(f"Error processing {file_name}: {e}")
            elif file_name.endswith(".pdf"):
                try:
                    self.process_pdf_document(file_path)
                except Exception as e:
                    print(f"Error processing {file_name}: {e}")

    def retrieve_similar_documents(self, query_text, top_n=3):
        """
        Retrieves documents from MongoDB that are most similar to the query text.
        """
        query_embedding = self.embed_text(query_text)
        all_documents = self.mongo_handler.retrieve_embeddings()

        similarities = []
        for doc in all_documents:
            title_similarity = cosine_similarity([query_embedding], [doc['title_embedding']])[0][0]
            body_similarities = [cosine_similarity([query_embedding], [body_emb])[0][0] for body_emb in doc.get('row_embeddings', [])]
            max_body_similarity = max(body_similarities) if body_similarities else 0
            total_similarity = 0.7 * title_similarity + 0.3 * max_body_similarity
            similarities.append((doc, total_similarity))

        similarities.sort(key=lambda x: x[1], reverse=True)
        return similarities[:top_n]


# Example usage
def usage():
    embedder = DocumentEmbedder()
    embedder.process_directory("Documents/Vie")

    query = "LIBERIS Compte"
    similar_docs = embedder.retrieve_similar_documents(query_text=query, top_n=1)

    for doc, similarity in similar_docs:
        print(f"File: {doc['file_name']}, Similarity: {similarity}")


if __name__ == "__main__":
    usage()
